from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_openai import ChatOpenAI
import os
os.environ["USER_AGENT"] = "MyApp/1.0"
from langchain_community.document_loaders import WebBaseLoader
from langchain_milvus import Milvus
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_huggingface import HuggingFaceEmbeddings
# from langchain.document_loaders import DocxLoader, PandasLoader
from dotenv import load_dotenv
# from langchain.document_loaders import DocxLoader, PandasLoader
# from langchain.document_loaders import PyPDFLoader  # 确保导入 PyPDFLoader
from langchain.prompts import PromptTemplate
from langchain_community.chat_models import ChatOllama
from langchain_core.output_parsers import JsonOutputParser, StrOutputParser
from langchain.prompts import PromptTemplate
from langchain.chains import GraphCypherQAChain
# from langchain_openai import ChatOpenAI
from langchain_core.output_parsers import StrOutputParser
from langchain_community.graphs import Neo4jGraph
from docx import Document
import pandas as pd
import pdfplumber
from langchain.schema import Document as LangchainDocument

os.environ["OPENAI_API_KEY"] = "sk-6TEh8i6lQB0Q9eVsF1gZ2QjPE9SC74mvhxPoMsU8AI4fdrmw"
os.environ["NEO4J_URI"] = "bolt://202.114.205.44:7687"
os.environ["NEO4J_USERNAME"] = "neo4j"
os.environ["NEO4J_PASSWORD"] = "12345678"
# load_dotenv()

neo4j_uri = os.getenv("NEO4J_URI")
neo4j_username = os.getenv("NEO4J_USERNAME")
neo4j_password = os.getenv("NEO4J_PASSWORD")
openai_api_key = os.getenv("OPENAI_API_KEY")

# 1. 加载文档
# 指定包含文档的目录
directory_path = 'C:/Users/wzw/work/规范'  # 替换为您的文档目录
docs = []

# os.environ["OPENAI_API_KEY"] = getpass.getpass()
# openai_api_key = openai_api_key

# 提供 API 密钥
llm = ChatOpenAI(temperature=0,model_name="gpt-4o")
# llm = ChatOpenAI(model="gpt-4")


# 创建图形实例
graph = Neo4jGraph(
    url=neo4j_uri,
    username=neo4j_username,
    password=neo4j_password
)

# 加载 .docx 文件
for filename in os.listdir(directory_path):
    if filename.endswith('.docx'):
        doc = Document(os.path.join(directory_path, filename))  # 使用正确的 Document 类
        text = '\n'.join(paragraph.text for paragraph in doc.paragraphs)
        docs.append(LangchainDocument(page_content=text, metadata={"title": filename, "url": filename}))

# 加载 .xls 和 .xlsx 文件
for filename in os.listdir(directory_path):
    if filename.endswith(('.xls', '.xlsx')):
        df = pd.read_excel(os.path.join(directory_path, filename))
        docs.append(LangchainDocument(page_content=df.to_string(), metadata={"title": filename, "url": filename}))

# 加载 .pdf 文件
for filename in os.listdir(directory_path):
    if filename.endswith('.pdf'):
        with pdfplumber.open(os.path.join(directory_path, filename)) as pdf:
            text = ''
            for page in pdf.pages:
                text += page.extract_text() + '\n'
            docs.append(LangchainDocument(page_content=text, metadata={"title": filename, "url": filename}))

# 输出加载的文档数量
print(f"Number of documents loaded: {len(docs)}")

# 使用 RecursiveCharacterTextSplitter 分割文档
text_splitter = RecursiveCharacterTextSplitter.from_tiktoken_encoder(
    chunk_size=500, chunk_overlap=50
)
doc_splits = text_splitter.create_documents(
    [doc.page_content for doc in docs], metadatas=[doc.metadata for doc in docs]
)

print(f"Number of chunks: {len(doc_splits)}")

# 将文档存储到 Milvus
vectorstore = Milvus.from_documents(
    documents=doc_splits,
    collection_name="rag_milvus",
    embedding=HuggingFaceEmbeddings(),
    connection_args={"host": "localhost", "port": "19530"},
)

# # 加载 .docx 文件
# for filename in os.listdir(directory_pa
#
# th):
#     if filename.endswith('.docx'):
#         doc = Document(os.path.join(directory_path, filename))
#         text = '\n'.join(paragraph.text for paragraph in doc.paragraphs)
#         docs.append({'content':text})  # 将加载的文档添加到列表中
#
# # 加载 .xls 文件
# for filename in os.listdir(directory_path):
#     if filename.endswith(('.xls', '.xlsx')):
#         df = pd.read_excel(os.path.join(directory_path, filename))
#         docs.append({'content',df.to_string()})  # 将加载的文档添加到列表中
#
# # 加载 .pdf 文件
# for filename in os.listdir(directory_path):
#     if filename.endswith('.pdf'):
#         with pdfplumber.open(os.path.join(directory_path, filename)) as pdf:
#             text = ''
#             for page in pdf.pages:
#                 text += page.extract_text() + '\n'
#             docs.append({'content':text})  # 将加载的文档添加到列表中
#
# # 输出加载的文档数量
# print(f"Number of documents loaded: {len(docs)}")
#
# # 生成向量表示
# embedding_model = HuggingFaceEmbeddings()
# vectors = [embedding_model.embed_query(doc.content) for doc in docs]  # 生成每个文档的向量
#
# # 将文档存储到 Milvus
# vectorstore = Milvus.from_documents(
#     documents=docs,
#     collection_name="my_documents",
#     embedding=HuggingFaceEmbeddings(),
#     connection_args={"uri": "./milvus_ingest.db"},
# )

# prompt = PromptTemplate(
#     template="""You are a grader assessing relevance
#     of a retrieved document to a user question. If the document contains keywords related to the user question,
#     grade it as relevant. It does not need to be a stringent test. The goal is to filter out erroneous retrievals.
#
#     Give a binary score 'yes' or 'no' score to indicate whether the document is relevant to the question.
#     Provide the binary score as a JSON with a single key 'score' and no premable or explaination.
#
#     Here is the retrieved document:
#     {document}
#
#     Here is the user question:
#     {question}
#     """,
#     input_variables=["question", "document"],
# )

cypher_prompt = PromptTemplate(
    template="""You are an expert at generating Cypher queries for Neo4j.
    Use the following schema to generate a Cypher query that answers the given question.
    Make the query flexible by using case-insensitive matching and partial string matching where appropriate.
    Focus on searching paper titles as they contain the most relevant information.

    Schema:
    {schema}

    Question: {question}

    Cypher Query:""",
    input_variables=["schema", "question"],
)

# QA prompt
qa_prompt = PromptTemplate(
    template="""You are an assistant for question-answering tasks. 
    Use the following Cypher query results to answer the question. If you don't know the answer, just say that you don't know. 
    Use three sentences maximum and keep the answer concise. If topic information is not available, focus on the paper titles.

    Question: {question} 
    Cypher Query: {query}
    Query Results: {context} 

    Answer:""",
    input_variables=["question", "query", "context"],
)



prompt_vector = PromptTemplate(
    template="""You are an assistant for question-answering tasks. 
    Use the following pieces of retrieved context to answer the question. If you don't know the answer, just say that you don't know. 
    Use three sentences maximum and keep the answer concise:
    Question: {question} 
    Context: {context} 
    Answer: 
    """,
    input_variables=["question", "document"],
)

question = "滑坡有哪些类型？"

#  milvus + llm
rag_chain = prompt_vector | llm | StrOutputParser()
retriever = vectorstore.as_retriever()

# generation = rag_chain.invoke({"context": docs, "question": question})

# neo4j + llm
graph_rag_chain = GraphCypherQAChain.from_llm(
    cypher_llm=llm,
    qa_llm=llm,
    validate_cypher=True,
    graph=Neo4jGraph,
    verbose=True,
    return_intermediate_steps=True,
    return_direct=True,
    cypher_prompt=cypher_prompt,
    qa_prompt=qa_prompt,
)

# 检索器得到的文档列表
docs = retriever.invoke(question)

# 查向量数据库得到的上下文
vector_context = rag_chain.invoke({"context": docs, "question": question})

# 查 neo4j得到的上下文
graph_context = graph_rag_chain.invoke({"query": question})

# 向量数据库 + neO4J的提示词
prompt = PromptTemplate(
    template="""You are an assistant for question-answering tasks. 
    Use the following pieces of retrieved context from a vector store and a graph database to answer the question. If you don't know the answer, just say that you don't know. 
    Use three sentences maximum and keep the answer concise:
    Question: {question} 
    Vector Context: {context} 
    Graph Context: {graph_context}
    Answer: 
    """,
    input_variables=["question", "context", "graph_context"],
)

composite_chain = prompt | llm | StrOutputParser()
# 基于之前两种数据库 综合查询
answer = composite_chain.invoke(
    {"question": question, "context": vector_context, "graph_context": graph_context}
)


print(answer)



















