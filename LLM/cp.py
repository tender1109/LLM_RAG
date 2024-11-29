from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_openai import ChatOpenAI
import os
os.environ["USER_AGENT"] = "MyApp/1.0"
from langchain_community.document_loaders import WebBaseLoader
from langchain_milvus import Milvus
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_huggingface import HuggingFaceEmbeddings
# from langchain.document_loaders import DocxLoader, PandasLoader
from dotenv import load_dotenv
# from langchain.document_loaders import DocxLoader, PandasLoader
from langchain_community.document_loaders import PyPDFLoader
from langchain.prompts import PromptTemplate
from langchain_community.chat_models import ChatOllama
from langchain_core.output_parsers import JsonOutputParser, StrOutputParser
from langchain.prompts import PromptTemplate
from langchain.chains import GraphCypherQAChain
# from langchain_openai import ChatOpenAI
from langchain_core.output_parsers import StrOutputParser
from langchain_community.graphs import Neo4jGraph
from docx import Document
import pandas as pd
import pdfplumber
from langchain import hub
from langchain.schema import Document as LangchainDocument
from langchain.prompts import PromptTemplate
from langchain.chains import GraphCypherQAChain
# from langchain_ollama import ChatOllama
from langchain_openai import ChatOpenAI
from langchain.vectorstores.milvus import Milvus

from langchain_core.embeddings import Embeddings
from common.chunk import text_to_chunk
from text2vec import SentenceModel

os.environ["OPENAI_API_KEY"] = "sk-6TEh8i6lQB0Q9eVsF1gZ2QjPE9SC74mvhxPoMsU8AI4fdrmw"
os.environ["NEO4J_URI"] = "bolt://202.114.205.44:7687"
os.environ["NEO4J_USERNAME"] = "neo4j"
os.environ["NEO4J_PASSWORD"] = "12345678"
# load_dotenv()

# # 设置 Clash 代理的地址和端口
# proxy_url = 'http://127.0.0.1'  # 或者使用 SOCKS 代理
# proxy_port = '7890'  # HTTP 代理端口
#
# # 设置环境变量
# os.environ['http_proxy'] = f'{proxy_url}:{proxy_port}'
# os.environ['https_proxy'] = f'{proxy_url}:{proxy_port}'



neo4j_uri = os.getenv("NEO4J_URI")
neo4j_username = os.getenv("NEO4J_USERNAME")
neo4j_password = os.getenv("NEO4J_PASSWORD")
openai_api_key = os.getenv("OPENAI_API_KEY")

# print(neo4j_uri,neo4j_username,neo4j_password,openai_api_key)

# 1. 加载文档
# 指定包含文档的目录
directory_path = 'C:/Users/wzw/work/规范'  # 替换为您的文档目录

docs = []

llm = ChatOpenAI(temperature=0,model_name="gpt-4o",base_url = "https://api.chatanywhere.tech/v1")
# llm = ChatOpenAI(model="gpt-4")


# 创建图形实例
graph = Neo4jGraph(url=neo4j_uri, username=neo4j_username, password=neo4j_password)

# 加载 .docx 文件
# for filename in os.listdir(directory_path):
#     if filename.endswith('.docx'):
#         doc = Document(os.path.join(directory_path, filename))
#         text = '\n'.join(paragraph.text for paragraph in doc.paragraphs)
#         docs.append(LangchainDocument(page_content=text, metadata={"title": filename, "url": filename}))

# 加载 .xls 和 .xlsx 文件
for filename in os.listdir(directory_path):
    if filename.endswith(('.xls', '.xlsx')):
        df = pd.read_excel(os.path.join(directory_path, filename))
        docs.append(LangchainDocument(page_content=df.to_string(), metadata={"title": filename, "url": filename}))

# # 加载 .pdf 文件
# for filename in os.listdir(directory_path):
#     if filename.endswith('.pdf'):
#         with pdfplumber.open(os.path.join(directory_path, filename)) as pdf:
#             text = ''
#             for page in pdf.pages:
#                 page_text = page.extract_text()
#                 if page_text:  # 处理 None
#                     text += page_text + '\n'
#             if text:  # 只有在有文本时才添加文档
#                 docs.append(LangchainDocument(page_content=text, metadata={"title": filename, "url": filename}))

# 输出加载的文档数量
print(f"Number of documents loaded: {len(docs)}")

loader = PyPDFLoader("C:/Users/wzw/work/规范/地质灾害风险识别技术规范.pdf")
#page是一个document对象
pages = loader.load_and_split()

full_text = "\n".join(page.page_content for page in pages)

chunks = text_to_chunk(full_text)
# 使用 RecursiveCharacterTextSplitter 分割文档
# text_splitter = RecursiveCharacterTextSplitter.from_tiktoken_encoder(
#     chunk_size=1000, chunk_overlap=20
# )
# chunks = text_to_chunk( [page.page_content for page in pages] )


# chunks = text_splitter.split_documents(pages)
# doc_splits = text_splitter.split_documents(
#     [doc.page_content for doc in docs]
# )

print(f"Number of chunks: {len(chunks)}")

model = SentenceModel('shibing624/text2vec-base-chinese')
embeddings = model.encode(chunks)

# 将文档存储到 Milvus
vectorstore = Milvus(
    embedding_function=model,
    connection_args={"host": "localhost", "port": "19530"},
    collection_name="rag_milvus",
).from_documents(
    chunks,
    embedding=embeddings,
)

cypher_prompt = PromptTemplate(
    template="""You are an expert at generating Cypher queries for Neo4j.
    Use the following schema to generate a Cypher query that answers the given question.
    Make the query flexible by using case-insensitive matching and partial string matching where appropriate.
    Focus on searching paper titles as they contain the most relevant information.

    Schema:
    {schema}

    Question: {question}

    Cypher Query:""",
    input_variables=["schema", "question"],
)

# QA prompt
qa_prompt = PromptTemplate(
    template="""You are an assistant for question-answering tasks. 
    Use the following Cypher query results to answer the question. If you don't know the answer, just say that you don't know. 
    Use three sentences maximum and keep the answer concise. If topic information is not available, focus on the paper titles.

    Question: {question} 
    Cypher Query: {query}
    Query Results: {context} 

    Answer:""",
    input_variables=["question", "query", "context"],
)



prompt_vector = PromptTemplate(
    template="""You are an assistant for question-answering tasks. 
    Use the following pieces of retrieved context to answer the question. If you don't know the answer, just say that you don't know. 
    Use three sentences maximum and keep the answer concise:
    Question: {question} 
    Context: {context} 
    Answer: 
    """,
    input_variables=["question", "context"],
)

question = "滑坡有哪些类型？"

#  milvus + llm
rag_chain = prompt_vector | llm | StrOutputParser()
retriever = vectorstore.as_retriever()

# generation = rag_chain.invoke({"context": docs, "question": question})

# neo4j + llm
graph_rag_chain = GraphCypherQAChain.from_llm(
    cypher_llm=llm,
    qa_llm=llm,
    validate_cypher=True,
    graph=graph,
    verbose=True,
    return_intermediate_steps=True,
    return_direct=True,
    cypher_prompt=cypher_prompt,
    qa_prompt=qa_prompt,
    allow_dangerous_requests=True,
)

# 检索器得到的文档列表
docs_qu = retriever.invoke(question)

# 查向量数据库得到的上下文
vector_context = rag_chain.invoke({"context": docs_qu, "question": question})

# 查 neo4j得到的上下文
graph_context = graph_rag_chain.invoke({"query": question})

# 向量数据库 + neO4J的提示词
prompt = PromptTemplate(
    template="""You are an assistant for question-answering tasks. 
    Use the following pieces of retrieved context from a vector store and a graph database to answer the question. If you don't know the answer, just say that you don't know. 
    Use three sentences maximum and keep the answer concise:
    Question: {question} 
    Vector Context: {context} 
    Graph Context: {graph_context}
    Answer: 
    """,
    input_variables=["question", "context", "graph_context"],
)

composite_chain = prompt | llm | StrOutputParser()
# 基于之前两种数据库 综合查询
answer = composite_chain.invoke(
    {"question": question, "context": vector_context, "graph_context": graph_context}
)


print(answer)
